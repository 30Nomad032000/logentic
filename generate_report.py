"""
Generate MCA Project Report for Hyper-Localized Multilingual Voice Assistant
Matching the format of the reference project report document.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PROJECT_TITLE = "HYPER-LOCALIZED MULTILINGUAL VOICE ASSISTANT USING EDGE COMPUTING AND AGENTIC WORKFLOWS"
STUDENT_NAME = "ALEXANDER SAJAN"
STUDENT_REG = "ASI24MCA-2006"
GUIDE_NAME = "Ms. Rintu Augustine"
GUIDE_TITLE = "Assistant professor"
DEPT = "DEPARTMENT OF COMPUTER APPLICATIONS"
COLLEGE = "Adi Shankara Institute of Engineering & Technology"
COLLEGE_FULL = "ADI SANKARA INSTITUTE OF ENGINEERING AND TECHNOLOGY VIDYA BHARATHI NAGAR, MATTOOR, KALADY"
COLLEGE_PLACE = "Kalady,683574"
YEAR = "2025-2026"
HOD_NAME = "Dr.Vincy Devi V K"


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_run(paragraph, text, font_name="Times New Roman", font_size=12,
            bold=False, italic=False, underline=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    # Ensure Times New Roman works for ASCII and East Asian
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)
    return run


def add_empty_paragraphs(doc, count, font_size=12):
    """Add empty paragraphs for spacing."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, 0, 0, 1.0)
        add_run(p, "", font_size=font_size)


def add_centered_text(doc, text, font_size=12, bold=False, spacing_before=0, spacing_after=0):
    """Add centered text paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, spacing_before, spacing_after, 1.0)
    add_run(p, text, font_size=font_size, bold=bold)
    return p


def add_body_text(doc, text, font_size=12, bold=False, alignment=None, spacing_after=6):
    """Add body text paragraph with justification."""
    p = doc.add_paragraph()
    p.alignment = alignment if alignment is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, 0, spacing_after, 1.5)
    add_run(p, text, font_size=font_size, bold=bold)
    return p


def add_heading_text(doc, text, font_size=16, bold=True):
    """Add a heading-style centered bold text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 12, 6, 1.0)
    add_run(p, text, font_size=font_size, bold=bold)
    return p


def add_section_break(doc):
    """Add a section break (new page)."""
    new_section = doc.add_section()
    new_section.top_margin = Cm(2.4)
    new_section.bottom_margin = Cm(0.5)
    new_section.left_margin = Cm(2.05)
    new_section.right_margin = Cm(2.05)
    new_section.page_width = Cm(21.0)
    new_section.page_height = Cm(29.7)
    return new_section


def create_cover_page(doc, include_guide=False):
    """Create cover/title page."""
    add_empty_paragraphs(doc, 3, font_size=12)

    # Title
    add_centered_text(doc, "A PROJECT REPORT", font_size=14, bold=True)
    add_centered_text(doc, "On", font_size=14)
    add_empty_paragraphs(doc, 1)
    add_centered_text(doc, PROJECT_TITLE, font_size=14, bold=True)

    # Submitted by
    add_centered_text(doc, "Submitted by", font_size=12, bold=True)
    add_centered_text(doc, f"                {STUDENT_NAME}", font_size=12, bold=True)
    add_centered_text(doc, f"                 ({STUDENT_REG})", font_size=12, bold=True)

    add_empty_paragraphs(doc, 3)

    add_centered_text(doc, "To", font_size=12, bold=True)
    add_empty_paragraphs(doc, 2)

    add_centered_text(doc, "The APJ Abdul Kalam Technological University", font_size=12, bold=True)
    add_centered_text(doc, "in partial fulfillment of the requirements for the award of the Degree of", font_size=12, bold=True)
    add_centered_text(doc, "Master of Computer Applications", font_size=12, bold=True)
    add_empty_paragraphs(doc, 1)

    if include_guide:
        add_centered_text(doc, "Guided by", font_size=14, bold=True)
        add_centered_text(doc, GUIDE_NAME, font_size=14, bold=True)
        add_centered_text(doc, GUIDE_TITLE, font_size=12)
        add_empty_paragraphs(doc, 9)
    else:
        add_empty_paragraphs(doc, 12)

    add_centered_text(doc, DEPT, font_size=12, bold=True)
    add_centered_text(doc, COLLEGE, font_size=14, bold=True)
    add_centered_text(doc, f" {COLLEGE_PLACE}", font_size=14, bold=True)
    add_centered_text(doc, YEAR, font_size=14, bold=True)


def create_declaration(doc):
    """Create declaration page."""
    add_heading_text(doc, "DECLARATION", font_size=16)
    add_empty_paragraphs(doc, 2)

    declaration_text = (
        f"I hereby declare that the project report entitled \u201c{PROJECT_TITLE}\u201d, "
        "submitted for partial fulfillment of the requirements for the award of the Degree of "
        "Master of Computer Applications of APJ Abdul Kalam Technological University, Kerala, "
        f"is a bonafide work carried out by me under the supervision of {GUIDE_NAME} {GUIDE_TITLE}. "
        "This submission represents my original work, and wherever the ideas or words of others "
        "have been included, they have been properly acknowledged and referenced. I further declare "
        "that I have adhered to the principles of academic honesty and integrity and have not "
        "misrepresented or fabricated any data, idea, fact, or source in this report. I understand "
        "that any violation of these principles may result in disciplinary action by the institute "
        "and/or the University and may also lead to legal consequences from the concerned sources. "
        "I also declare that this report has not previously formed the basis for the award of any "
        "degree, diploma, or similar title of any other University or Institution."
    )
    add_body_text(doc, declaration_text, font_size=12, spacing_after=0)

    add_empty_paragraphs(doc, 5)

    # Place and signature
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    add_run(p, "Place :KALADY\tSignature", font_size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    add_run(p, "Date :", font_size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    add_run(p, STUDENT_NAME, font_size=10)


def create_certificate(doc):
    """Create certificate page."""
    # College header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 0, 1.0)
    add_run(p, f"{COLLEGE_FULL} {DEPT}", font_size=14, bold=True)

    add_empty_paragraphs(doc, 3)

    add_heading_text(doc, "CERTIFICATE", font_size=16)

    cert_text = (
        f"This is to certify that the project report entitled \u201c{PROJECT_TITLE}\u201d "
        f"submitted by {STUDENT_NAME} ({STUDENT_REG}) to the APJ Abdul Kalam Technological University "
        "in partial fulfillment of the requirements for the award of the Degree of Master of Computer "
        "Applications is a bonafide record of the project work carried out by him under my guidance "
        f"and supervision during the academic year {YEAR}. "
    )
    add_body_text(doc, cert_text, font_size=12, spacing_after=0)

    add_empty_paragraphs(doc, 3)

    # Guide and HOD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    add_run(p, f"Asst. Prof. {GUIDE_NAME.replace('Ms. ', '')}\t{HOD_NAME}", font_size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    add_run(p, "Master of Computer Applications\tHead of the Department", font_size=12)

    add_empty_paragraphs(doc, 3)

    add_body_text(doc, "Submitted to presentation and evaluation on............at\u2026\u2026\u2026\u2026", font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, spacing_after=0)

    add_empty_paragraphs(doc, 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    add_run(p, "Examiner 1                                                                                          Examiner2", font_size=12, bold=True)

    add_empty_paragraphs(doc, 3)

    add_centered_text(doc, "College Seal", font_size=12, bold=True)


def create_acknowledgement(doc):
    """Create acknowledgement page."""
    add_heading_text(doc, "ACKNOWLEDGEMENT", font_size=16)
    add_empty_paragraphs(doc, 2)

    paras = [
        "I am profoundly grateful for the support and guidance that have been instrumental in the successful completion of this project.",

        f"First and foremost, I extend my heartfelt thanks to Asst. Prof. {GUIDE_NAME.replace('Ms. ', '')}, "
        "Department of Computer Applications, whose expertise, mentorship, and constructive feedback have been "
        "invaluable throughout the project. Her insights not only refined my technical skills but also helped "
        "shape this project into a comprehensive solution for building a hyper-localized multilingual voice "
        "assistant using edge computing and agentic workflows.",

        f"I am equally thankful to {COLLEGE} for providing the resources, infrastructure, and learning "
        "environment that enabled me to delve deeply into advanced machine learning, natural language processing, "
        "and edge computing. The access to advanced tools, laboratories, and a supportive academic community "
        "has significantly enriched my experience and knowledge in AI-driven software engineering and modern "
        "voice-based systems.",

        "I would like to acknowledge the open-source community and various online forums for their rich "
        "repositories of knowledge, tools, and technical discussions, which served as critical references "
        "throughout the development process. The availability of modern frameworks such as PyTorch, "
        "Transformers, LangGraph, and FastAPI has been fundamental to the successful implementation of this project.",

        "I am also grateful to all the individuals who provided valuable feedback during the testing phase "
        "of the application, helping me identify areas for improvement and ensuring the system meets "
        "real-world user requirements across multiple Indian languages.",

        "Finally, I express my sincere gratitude to my family and friends for their unwavering support, "
        "encouragement, and patience throughout this journey. Their belief in my abilities has been a "
        "constant source of motivation."
    ]

    for para_text in paras:
        add_body_text(doc, para_text, font_size=12)
        add_empty_paragraphs(doc, 1)


def create_abstract(doc):
    """Create abstract page."""
    add_heading_text(doc, "ABSTRACT", font_size=16)
    add_empty_paragraphs(doc, 1)

    paras = [
        (
            f"This project proposes an intelligent system for \u201c{PROJECT_TITLE}\u201d, "
            "designed to enable seamless interaction with digital services through natural voice "
            "commands in regional Indian languages. The primary problem addressed is the digital "
            "divide faced by millions of users in India who are unable to effectively interact with "
            "technology due to limited support for their native languages, particularly in voice-based "
            "interfaces. This challenge is especially pronounced in rural and semi-urban areas where "
            "English proficiency is low and existing voice assistants fail to serve users adequately."
        ),
        (
            "The system\u2019s core objective is to bridge this accessibility gap by providing a "
            "voice assistant that supports 11 Indian languages including Malayalam, Hindi, Tamil, "
            "Telugu, Bengali, Marathi, Gujarati, Kannada, Punjabi, Urdu, and English. The solution "
            "integrates a complete processing pipeline comprising Automatic Speech Recognition (ASR) "
            "using the Pingala V1 model for CPU-optimized multilingual transcription, bidirectional "
            "translation using IndicTrans2 for converting between Indian languages and English, "
            "intent detection using keyword-based and ML-based classifiers, response generation "
            "using locally-deployed Qwen3 large language models via the llama.cpp backend, and "
            "Text-to-Speech (TTS) synthesis using Meta MMS-TTS for offline multilingual audio output."
        ),
        (
            "The system is structured using a modular architecture with LangGraph-based agentic "
            "workflows that intelligently route user queries to specialized agents\u2014including "
            "Information, Task Management, Conversation, and Smart Home agents\u2014enabling "
            "context-aware and domain-specific responses. The entire system is designed for edge "
            "deployment on Raspberry Pi 5 hardware, utilizing quantized models and CPU-first "
            "inference strategies to achieve low-latency, privacy-preserving, and offline-capable "
            "operation. A FastAPI-based REST and WebSocket API serves as the backend, complemented "
            "by a React-based monitoring dashboard for real-time system analytics. The project "
            "demonstrates a sophisticated approach to building accessible voice interfaces for "
            "under-resourced languages while maintaining deployment flexibility across edge and "
            "cloud scenarios."
        ),
    ]

    for para_text in paras:
        add_body_text(doc, para_text, font_size=12)
        add_empty_paragraphs(doc, 1)


def main():
    doc = Document()

    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    # Set page margins for first section (A4)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(2.05)
    section.right_margin = Cm(2.05)

    # === Page 1: Cover page without guide ===
    create_cover_page(doc, include_guide=False)

    # === Page 2: Cover page with guide ===
    add_section_break(doc)
    create_cover_page(doc, include_guide=True)

    # === Page 3: Declaration ===
    add_section_break(doc)
    create_declaration(doc)

    # === Page 4: Certificate ===
    add_section_break(doc)
    create_certificate(doc)

    # === Page 5: Acknowledgement ===
    add_section_break(doc)
    create_acknowledgement(doc)

    # === Page 6: Abstract ===
    add_section_break(doc)
    create_abstract(doc)

    # Save
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Project_Report_Voice_Assistant.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    main()
