"""
Generate the complete Project Report as a .docx document.
Embeds matplotlib figures and Puppeteer screenshots.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ASSETS = Path(__file__).parent / "report_assets"
OUT = Path(__file__).parent / "Project_Report_Voice_Assistant_v2.docx"

doc = Document()

# ══════════════════════════════════════════════════════════════
# GLOBAL STYLES
# ══════════════════════════════════════════════════════════════

# -- Section / page setup (A4) --
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.75)
section.right_margin = Cm(2.86)

# -- Header --
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.text = "Hyper-Localized Multilingual Voice Assistant"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "Times New Roman"

# -- Footer --
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_f = fp.add_run("Department of Computer Applications – ASIET")
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run_f.font.name = "Times New Roman"

# -- Update Normal style --
style_normal = doc.styles["Normal"]
style_normal.font.name = "Times New Roman"
style_normal.font.size = Pt(12)
pf = style_normal.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(4)

# ── Helper functions ──────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set table cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading1(text):
    """Chapter heading: 14pt, bold, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    # Add page break before chapter (except first)
    if text != "1. INTRODUCTION":
        p.paragraph_format.page_break_before = True
    return p


def add_heading2(text):
    """Section heading: 12pt, bold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.85)
    return p


def add_heading3(text):
    """Subsection heading: 12pt, bold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.85)
    return p


def add_heading4(text):
    """Italic heading in Cambria."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = "Cambria"
    run.font.color.rgb = RGBColor(0x36, 0x5F, 0x91)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(text):
    """Justified body paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


def add_body_bold_start(bold_text, rest_text):
    """Paragraph starting with bold text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(rest_text)
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"
    return p


def add_bullet(text, bold_prefix=None):
    """Bulleted list item."""
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = "Times New Roman"
        r2 = p.add_run(text)
        r2.font.size = Pt(12)
        r2.font.name = "Times New Roman"
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
    return p


def add_figure(filename, caption, width=Inches(5.5)):
    """Add an image with a centered caption."""
    fpath = ASSETS / filename
    if fpath.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fpath), width=width)
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    else:
        add_body(f"[Figure: {caption} — image not found: {filename}]")


def add_table(headers, rows):
    """Add a bordered table with shaded header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D6E4F0")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Times New Roman"

    doc.add_paragraph()  # spacer
    return table


def add_code_block(code, purpose=None):
    """Add a code block with gray background and monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # Set paragraph shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    if purpose:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = pp.add_run(f"Purpose: {purpose}")
        r.italic = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ══════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════

add_heading1("1. INTRODUCTION")

add_heading2("1.1 Brief Introduction About the Project")

add_body(
    'The project "Hyper-Localized Multilingual Voice Assistant" is an innovative system designed to enable '
    "seamless interaction with digital services using natural voice commands in regional Indian languages. "
    "By leveraging modern edge computing (Raspberry Pi 5) and agentic AI workflows (LangGraph), this project "
    "aims to bridge the digital divide for millions of users who are more comfortable communicating in their "
    "native tongue. The system transforms voice commands in languages like Malayalam, Hindi, Tamil, Telugu, "
    "and 7 other Indian languages into actionable responses \u2014 performing tasks such as information retrieval, "
    "task management, smart home control, and general conversation."
)

add_body(
    "The conventional approach to voice assistants is dominated by English-centric, cloud-dependent solutions "
    "like Alexa, Siri, and Google Assistant. These systems offer limited or poor support for regional Indian "
    "languages, require constant internet connectivity, and often fail to understand the nuances of local "
    "dialects and accents. This creates a significant accessibility barrier for the vast majority of India\u2019s "
    "population, who speak languages other than English as their primary language. Furthermore, the reliance "
    "on cloud processing introduces privacy concerns and latency issues, particularly in rural and semi-urban "
    "areas with unreliable internet connectivity."
)

add_body(
    "This voice assistant provides a unified, modular platform that seamlessly connects users with AI-powered "
    "services through voice in their native language. The system is fundamentally broken down into two primary "
    "interfaces:"
)

add_body_bold_start(
    "The Edge Device (Raspberry Pi 5): ",
    "The hardware component captures audio input through a microphone, processes wake-word detection locally, "
    "and communicates with the backend server. It runs optimized, quantized AI models (Qwen 2.5 via llama.cpp) "
    "for low-latency inference directly on-device, ensuring the system works even with limited internet "
    "connectivity."
)

add_body_bold_start(
    "The Dashboard/Management Interface: ",
    "A React-based real-time monitoring dashboard provides administrators and developers with complete "
    "visibility into the system\u2019s operation. Features include real-time conversation logs, component health "
    "monitoring (ASR, LLM, TTS status), latency breakdown analysis, language distribution statistics, and "
    "an interactive \u201cTry It\u201d panel for testing voice interactions directly from the browser."
)

add_body(
    "The system\u2019s design prioritizes a smooth pipeline flow from audio input to audio response: "
    "Audio \u2192 ASR (Whisper/MMS) \u2192 Translation (ML\u2192EN via IndicTrans2) \u2192 Intent Detection \u2192 "
    "LLM Response (Qwen 2.5) \u2192 Translation (EN\u2192ML) \u2192 TTS (MMS-TTS) \u2192 Audio Response. "
    "Each step is independently configurable, cacheable, and optimized for edge deployment."
)

add_heading2("1.2 Project Background")

add_body(
    "India\u2019s linguistic diversity \u2014 with 22 officially recognized languages and over 19,500 dialects \u2014 "
    "presents a unique challenge for voice technology. According to the 2011 Census, only about 10% of India\u2019s "
    "population speaks English, yet the majority of voice assistant technology is optimized primarily for English. "
    "This creates a massive digital accessibility gap, particularly in rural and semi-urban areas where regional "
    "languages dominate daily communication."
)

add_body(
    "The primary motivation behind the development of this Hyper-Localized Multilingual Voice Assistant is to "
    "directly address this language barrier by building a voice-first interface that natively supports Indian "
    "languages. The project is conceived as a comprehensive AI ecosystem that operates at the edge. By utilizing "
    "offline-capable models running on a Raspberry Pi 5, the system ensures accessibility even in areas with poor "
    "or no internet connectivity. The core objective is to transition the user experience from an English-dependent, "
    "cloud-only model to a multilingual, edge-first model that respects linguistic diversity."
)

add_body(
    "Furthermore, the project provides developers and researchers with a monitoring dashboard and a modular "
    "pipeline architecture. The LangGraph-based agent orchestrator enables intelligent routing of user queries "
    "to specialized agents (information, task management, smart home, conversation), while the comprehensive "
    "testing suite and benchmarking tools ensure reliability and performance optimization."
)

add_heading2("1.3 Existing System")

add_body(
    "The existing landscape of voice assistants \u2014 dominated by Amazon Alexa, Google Assistant, Apple Siri, "
    "and Microsoft Cortana \u2014 suffers from several critical limitations when it comes to serving India\u2019s "
    "linguistically diverse population:"
)

add_bullet("While some progress has been made, these systems support only a handful of Indian languages, and "
           "even then, the accuracy and naturalness of interactions are significantly below par compared to English.",
           bold_prefix="Limited Indian Language Support: ")
add_bullet("All major voice assistants require constant, high-bandwidth internet connectivity for processing. "
           "In rural India, where internet penetration is approximately 31%, this is a fundamental barrier to adoption.",
           bold_prefix="Cloud Dependency: ")
add_bullet("All audio data is processed on remote cloud servers, raising significant privacy concerns \u2014 "
           "especially for users in healthcare, government, and education sectors.",
           bold_prefix="Privacy Concerns: ")
add_bullet("None of the mainstream voice assistants offer the ability to run entirely on local hardware, "
           "limiting deployment options in bandwidth-constrained environments.",
           bold_prefix="No Edge Computing: ")
add_bullet("Traditional voice assistants use simple intent-response mapping without sophisticated multi-agent "
           "reasoning capabilities.",
           bold_prefix="Lack of Agentic Workflows: ")
add_bullet("The round-trip to cloud servers introduces 500ms\u20132000ms of latency, which significantly "
           "degrades the conversational experience.",
           bold_prefix="High Latency: ")

add_heading2("1.4 Proposed System")

add_body(
    "The proposed Hyper-Localized Multilingual Voice Assistant addresses every limitation of the existing "
    "systems through a comprehensive, modular architecture:"
)

add_bullet("Malayalam (primary), Hindi, Tamil, Telugu, Bengali, Marathi, Gujarati, Kannada, Punjabi, Urdu, "
           "and English using state-of-the-art IndicTrans2 translation and MMS-TTS synthesis.",
           bold_prefix="Full Support for 11 Indian Languages: ")
add_bullet("Core AI models (Qwen 2.5 for LLM, MMS for ASR/TTS, IndicTrans2 for translation) are quantized "
           "and optimized to run on a Raspberry Pi 5, enabling offline or low-connectivity operation.",
           bold_prefix="Edge-First Design: ")
add_bullet("LangGraph-based orchestrator with specialized agents (Info, Task, Chat, Smart Home) that "
           "intelligently route and handle diverse user requests.",
           bold_prefix="Agentic AI Workflows: ")
add_bullet("Audio processing happens locally on the edge device; no data leaves the local network unless "
           "explicitly configured.",
           bold_prefix="Privacy by Design: ")
add_bullet("React + TypeScript dashboard with WebSocket streaming provides real-time visibility into system "
           "health, conversation logs, and performance metrics.",
           bold_prefix="Real-time Monitoring Dashboard: ")
add_bullet("Each component (ASR, Translation, LLM, TTS) can be independently swapped, configured, and "
           "optimized through a YAML configuration file.",
           bold_prefix="Modular, Configurable Pipeline: ")

add_heading2("1.5 Features and Limitations")

add_body("Features:")
features = [
    "Multilingual voice interaction in 11 Indian languages",
    "Edge computing deployment on Raspberry Pi 5",
    "LangGraph-based multi-agent system with 4 specialized agents",
    "Real-time monitoring dashboard with WebSocket streaming",
    "Multiple ASR backends (Whisper, MMS, Pingala, Indic-Whisper)",
    "Multiple TTS backends (MMS-TTS, Cartesia, Indic Parler-TTS, Chatterbox)",
    "Audio denoising and preprocessing pipeline",
    "Task management (reminders, alarms, to-do lists)",
    "Web search integration for real-time information",
    "Response caching with LRU + TTL strategies",
    "Comprehensive REST API and WebSocket endpoints",
    "Docker containerization for easy deployment",
    "Benchmarking and performance profiling tools",
]
for f in features:
    add_bullet(f)

add_body("Limitations:")
limitations = [
    "Initial model download requires internet connectivity (~2\u20134 GB for all models)",
    "Edge device (RPi 5) has limited memory (8 GB), restricting larger model sizes",
    "Translation accuracy varies across language pairs; Malayalam has the highest accuracy",
    "Real-time TTS synthesis on CPU can introduce 200\u2013500 ms latency",
    "Wake word detection is basic keyword-based, not deep learning-based",
    "No multi-turn conversation memory persistence across sessions currently",
]
for l in limitations:
    add_bullet(l)


# ══════════════════════════════════════════════════════════════
# CHAPTER 2 — REQUIREMENTS AND ANALYSIS
# ══════════════════════════════════════════════════════════════

add_heading1("2. REQUIREMENTS AND ANALYSIS")

add_heading2("2.1 Hardware Requirements")

add_table(
    ["Category", "Specification"],
    [
        ["Processor", "Raspberry Pi 5 (Broadcom BCM2712, 2.4 GHz quad-core Cortex-A76) or Intel Core i3+ for development"],
        ["RAM", "8 GB (RPi 5) / 16 GB+ (development machine)"],
        ["Storage", "32 GB microSD (RPi 5) / 500 GB SSD (development)"],
        ["GPU", "NVIDIA CUDA-capable GPU recommended for training/fast inference"],
        ["Operating System", "Raspberry Pi OS (Bookworm) / Ubuntu 22.04 / Windows 11"],
        ["Audio Hardware", "USB microphone + 3.5 mm speaker (RPi) / Built-in audio (dev)"],
        ["Network", "Wi-Fi / Ethernet for initial setup; offline operation supported"],
        ["Programming Language", "Python 3.10+"],
    ],
)

add_heading2("2.2 Software Requirements")

add_table(
    ["Requirement", "Details"],
    [
        ["Python Runtime", "Python 3.10+ with pip"],
        ["ML Framework", "PyTorch 2.x with CUDA support (optional)"],
        ["ASR Engine", "OpenAI Whisper / Meta MMS / Pingala V1"],
        ["LLM", "Qwen 2.5 Instruct (0.5B\u20137B variants) via llama.cpp"],
        ["Translation", "IndicTrans2 (AI4Bharat) \u2014 distributed (~211M params)"],
        ["TTS", "MMS-TTS (Meta) / Cartesia API / Indic Parler-TTS"],
        ["Agent Framework", "LangGraph + LangChain"],
        ["Web Framework", "FastAPI + Uvicorn"],
        ["Frontend", "React 18 + TypeScript + Vite + TailwindCSS"],
        ["Database", "SQLite with WAL mode"],
        ["Containerization", "Docker + Docker Compose"],
    ],
)


# ══════════════════════════════════════════════════════════════
# CHAPTER 3 — SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════

add_heading1("3. SYSTEM DESIGN")

add_heading2("3.1 Architecture")

add_body(
    "The system follows a three-tier architecture optimized for edge deployment, real-time processing, "
    "and administrative monitoring. The pipeline processes voice commands through a chain of AI models, "
    "each independently configurable and cacheable."
)

add_figure("pipeline_flow.png", "Fig 3.1 \u2014 End-to-End Voice Processing Pipeline", Inches(5.8))

add_body(
    "The architecture is organized into three tiers:"
)

add_body_bold_start("1. Edge Tier (Raspberry Pi 5): ",
    "Captures audio input via USB microphone, performs local wake-word detection, runs quantized models "
    "(Qwen 2.5 Q8_0 via llama.cpp), and communicates with the processing tier over HTTP/WebSocket.")
add_body_bold_start("2. Processing Tier (FastAPI Server): ",
    "Hosts the modular AI pipeline \u2014 ASR, Translation, LLM, TTS \u2014 behind a FastAPI REST API with "
    "WebSocket streaming support. Includes the LangGraph agent orchestrator, SQLite database, response "
    "cache, and XiaoZhi protocol bridge for EdgeX/ESP32 devices.")
add_body_bold_start("3. Presentation Tier (React Dashboard): ",
    "A React 18 + TypeScript single-page application providing real-time monitoring: conversation logs, "
    "component health, latency breakdown, language distribution, and an interactive \u201cTry It\u201d panel.")

add_figure("architecture.png", "Fig 3.1a \u2014 Three-Tier System Architecture", Inches(5))

add_heading2("3.2 UML Diagram")

add_figure("uml_class.png", "Fig 3.2 \u2014 UML Class Diagram of Core Components", Inches(5.8))

add_heading2("3.3 Use-Case Diagram")

add_figure("usecase.png", "Fig 3.3 \u2014 Use-Case Diagram", Inches(5.2))


# ══════════════════════════════════════════════════════════════
# CHAPTER 4 — AGILE TECHNOLOGY OVERVIEW
# ══════════════════════════════════════════════════════════════

add_heading1("4. AGILE TECHNOLOGY OVERVIEW")

add_body(
    "Agile technology refers to a set of principles and practices for software development that prioritize "
    "flexibility, collaboration, and rapid delivery. It emphasizes iterative development, continuous user "
    "feedback, and adaptability to changing requirements, rather than following a strict, linear process "
    "like traditional waterfall methodologies."
)

add_body_bold_start("Customer Collaboration Over Contract Negotiation: ",
    "Agile emphasizes working closely with stakeholders, ensuring that the voice assistant meets the "
    "evolving needs of users across diverse linguistic communities.")
add_body_bold_start("Responding to Change Over Following a Plan: ",
    "Agile allows the team to incorporate feedback, add features such as new ASR backends, additional "
    "language support, or task management capabilities during development.")
add_body_bold_start("Delivering Working Software Frequently: ",
    "Agile teams deliver functional increments at regular intervals \u2014 such as the core pipeline, "
    "dashboard interface, Docker deployment, and testing suite.")

add_heading3("4.1 Introduction to Scrum")

add_body(
    "Scrum is one of the most widely adopted frameworks within Agile technology. It provides a structured "
    "way to organize teams, tasks, and product development. The Scrum framework is based on fixed-length "
    "iterations called Sprints, which typically last from 1 to 4 weeks. Each Sprint results in a potentially "
    "shippable product increment."
)

add_body("Roles in Scrum:")
add_bullet("Represents the customer or stakeholders and ensures the team works on the most valuable features.",
           bold_prefix="Product Owner: ")
add_bullet("Facilitates the Scrum process, removes impediments, and ensures adherence to Scrum principles.",
           bold_prefix="Scrum Master: ")
add_bullet("A cross-functional group who work together to complete the Sprint Backlog.",
           bold_prefix="Development Team: ")

add_body("Scrum Ceremonies:")
add_bullet("The team decides what work will be done in the Sprint.", bold_prefix="Sprint Planning: ")
add_bullet("A daily meeting to discuss progress, plans, and roadblocks.", bold_prefix="Daily Standup: ")
add_bullet("The team demonstrates the increment to stakeholders.", bold_prefix="Sprint Review: ")
add_bullet("The team reflects on the Sprint and discusses improvements.", bold_prefix="Sprint Retrospective: ")

add_heading3("4.2 Principles or Methodology Used in Scrum")

add_bullet("Team members commit to the Sprint goals and backlog items.", bold_prefix="Commitment: ")
add_bullet("Teams focus on a small number of tasks per Sprint.", bold_prefix="Focus: ")
add_bullet("Transparency about progress and challenges.", bold_prefix="Openness: ")
add_bullet("Valuing the contributions of all members.", bold_prefix="Respect: ")
add_bullet("Empowered to speak up about potential issues.", bold_prefix="Courage: ")
add_bullet("Each Sprint provides an opportunity to improve product and processes.", bold_prefix="Continuous Improvement: ")

add_heading3("4.4 Product Backlog")

add_table(
    ["Sl.No", "User Story", "Est. Hours", "Priority"],
    [
        ["1", "As a user, I want to speak a command in Malayalam so the system understands and responds in my language.", "30", "High"],
        ["2", "As a developer, I want a modular ASR pipeline so I can swap between Whisper, MMS, and Pingala engines.", "24", "High"],
        ["3", "As a user, I want the system to intelligently route my queries to specialized agents (info, task, chat).", "30", "High"],
        ["4", "As an admin, I want a real-time dashboard to monitor system health, conversations, and latency.", "26", "High"],
        ["5", "As a user, I want to set reminders and manage tasks through voice commands.", "20", "High"],
        ["6", "As a developer, I want Docker deployment so the system can be deployed on any edge device easily.", "22", "High"],
    ],
)

add_heading3("4.5 Sprint Planner")

# Sprint Backlogs 1-6
sprint_data = [
    ("Sprint Backlog \u2013 1", "Project Initialization & Core Pipeline",
     [["1.1", "Set up project structure, CLAUDE.md, Python environment, core dependencies.", "Establish foundational AI pipeline.", "10", "Project runs with uvicorn."],
      ["1.2", "Implement Whisper ASR, MMS-TTS, Qwen LLM, IndicTrans2 modules with FastAPI.", "", "30", "End-to-end pipeline processes input."]]),
    ("Sprint Backlog \u2013 2", "Frontend Dashboard & UI",
     [["2.1", "Create React frontend with Vite + TypeScript scaffolding.", "Provide monitoring and interaction interface.", "14", "React app builds at localhost:3000."],
      ["2.2", "Migrate to TailwindCSS, implement dashboard components.", "", "16", "Dashboard displays live data."],
      ["2.3", "Add voice visualizer and WebSocket streaming.", "", "10", "Real-time audio interaction works."]]),
    ("Sprint Backlog \u2013 3", "Docker, Optimization & Documentation",
     [["3.1", "Create multi-stage Dockerfile and docker-compose.yml.", "Production-readiness and optimization.", "15", "System runs in Docker."],
      ["3.2", "Implement LRU cache, model optimizer, benchmarking.", "", "25", "Cache tracking works."]]),
    ("Sprint Backlog \u2013 4", "Testing, EdgeX Bridge & Agent Enhancement",
     [["4.1", "Add comprehensive test suite (test_basic, test_api, test_database, etc.).", "Quality assurance and edge integration.", "25", "All tests pass with pytest."],
      ["4.2", "Implement XiaoZhi protocol bridge and LLM intent classification.", "", "15", "EdgeX devices connect."]]),
    ("Sprint Backlog \u2013 5", "Agent Routing, Seed Data & Demo",
     [["5.1", "Add agent routing display in dashboard and seed demo data.", "Demo-ready system.", "18", "Dashboard shows agent routing."],
      ["5.2", "Switch to offline-only models for reliable demo.", "", "12", "Full demo works offline."]]),
    ("Sprint Backlog \u2013 6", "Task Management, Web Search & Audio",
     [["6.1", "Add task management with Task model, API endpoints, TaskPanel UI.", "Feature-complete system.", "20", "Users can manage tasks."],
      ["6.2", "Implement web search and audio denoising pipeline.", "", "20", "Search results appear; noisy audio cleaned."]]),
]

for title, backlog_item, rows in sprint_data:
    add_heading3(title)
    headers = ["Sprint ID", "Sprint Task", "Sprint Goal", "Est. Hours", "Acceptance Criteria"]
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "D6E4F0")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
        # Merge sprint goal if empty
        if ri > 0 and not rows[ri][2]:
            table.rows[1].cells[2].merge(table.rows[ri + 1].cells[2])
    doc.add_paragraph()

add_heading3("4.6 Ideal Burndown Chart")

add_figure("burndown.png", "Fig 4.1 \u2014 Sprint Burndown Chart (230 Total Hours)", Inches(5))

add_table(
    ["Sprint", "Planned Hours", "Actual Completed", "Remaining"],
    [
        ["Start", "\u2014", "\u2014", "230"],
        ["Sprint 1", "40", "42", "188"],
        ["Sprint 2", "40", "38", "150"],
        ["Sprint 3", "40", "40", "110"],
        ["Sprint 4", "40", "43", "67"],
        ["Sprint 5", "30", "30", "37"],
        ["Sprint 6", "40", "37", "0"],
    ],
)


# ══════════════════════════════════════════════════════════════
# CHAPTER 5 — SYSTEM DESIGN (TECH)
# ══════════════════════════════════════════════════════════════

add_heading1("5. SYSTEM DESIGN")

add_heading2("5.1 Technology / Framework")

add_body("Front-End and Back-End Technologies:")
add_table(
    ["Category", "Technology", "Version"],
    [
        ["Frontend", "React, TypeScript, TailwindCSS", "React 18.3, TS 5.6, Tailwind 3.4"],
        ["Backend", "FastAPI, Uvicorn, Python", "FastAPI 0.100+, Python 3.10+"],
        ["ML Framework", "PyTorch, Transformers, Accelerate", "PyTorch 2.x"],
        ["ASR", "OpenAI Whisper, Meta MMS, Pingala V1", "Latest stable"],
        ["TTS", "MMS-TTS, Cartesia, Indic Parler-TTS, Chatterbox", "Latest stable"],
        ["Translation", "IndicTrans2 (AI4Bharat), SentencePiece", "~211M params"],
        ["LLM", "Qwen 2.5 Instruct via llama.cpp", "0.5B\u20137B variants"],
        ["Agents", "LangGraph, LangChain", "Latest stable"],
        ["Database", "SQLite with WAL mode", "3.x"],
        ["Containerization", "Docker, Docker Compose", "Latest stable"],
    ],
)

add_body("Versioning and Software Tools:")
add_table(
    ["Purpose", "Tool", "Version"],
    [
        ["Version Control", "Git + GitHub", "Latest stable"],
        ["Code Editor", "Visual Studio Code", "Latest stable"],
        ["Package Manager", "pip, npm", "Latest stable"],
        ["Code Formatter", "Black (Python), Prettier (TS)", "Latest stable"],
        ["Linter", "Ruff (Python), ESLint (TS)", "Latest stable"],
        ["Testing", "pytest, pytest-asyncio", "Latest stable"],
    ],
)

add_figure("languages.png", "Fig 5.1 \u2014 Supported Language Distribution", Inches(4))
add_figure("latency.png", "Fig 5.2 \u2014 Component Latency Breakdown (GPU vs CPU)", Inches(5))
add_figure("memory.png", "Fig 5.3 \u2014 Model Memory Footprint", Inches(5))

add_heading2("5.2 Module Description")

add_table(
    ["Module Name", "Description"],
    [
        ["ASR Module (asr/)", "Speech-to-text conversion with multiple engines: OpenAI Whisper, Meta MMS, Pingala V1, Indic-Whisper. Includes audio denoising."],
        ["TTS Module (tts/)", "Text-to-speech supporting MMS-TTS (offline), Cartesia (cloud), Indic Parler-TTS, Chatterbox. A/B testing capability."],
        ["Translation (translation/)", "Bidirectional ML\u2194EN using IndicTrans2 from AI4Bharat. 11 Indian languages, auto language detection."],
        ["LLM Module (llm/)", "Response generation using Qwen 2.5 Instruct (0.5B\u20137B). Transformers and llama.cpp backends."],
        ["Agent Orchestrator (agents/)", "LangGraph workflow engine routing to: Info Agent, Task Agent, Chat Agent, Smart Home Agent."],
        ["Pipeline (pipeline/)", "End-to-end orchestrator: ASR\u2192Translation\u2192LLM\u2192TTS with timing metrics. Legacy and Pipecat modes."],
        ["API Module (api/)", "FastAPI REST endpoints and WebSocket streaming. XiaoZhi protocol bridge for EdgeX/ESP32."],
        ["Database (db/)", "SQLite persistence for conversations, messages, tasks, metrics. WAL mode for concurrency."],
        ["Cache (cache/)", "LRU cache with TTL for translation, TTS, LLM responses. Hit rate tracking and eviction stats."],
        ["Edge Module (edge/)", "Raspberry Pi 5 client with audio I/O, wake word detection, async HTTP/WebSocket."],
    ],
)


# ══════════════════════════════════════════════════════════════
# CHAPTER 6 — CODING
# ══════════════════════════════════════════════════════════════

add_heading1("6. CODING")

add_heading2("6.1 Coding Standards Used in the Project")

add_body("I. Python Language Conventions")
add_bullet("PEP 8 compliant: snake_case for functions/variables, PascalCase for classes, UPPER_SNAKE_CASE for constants.", bold_prefix="Naming: ")
add_bullet("Extensive use of typing module \u2014 TypedDict, Optional, List, Dict \u2014 and dataclass decorators.", bold_prefix="Type Hints: ")
add_bullet("Triple-quote docstrings on all modules, classes, and public methods.", bold_prefix="Documentation: ")
add_bullet("Black formatter (88 char line length), Ruff linter for error catching.", bold_prefix="Formatting: ")
add_bullet("async/await for FastAPI endpoints, WebSocket handlers, and edge communication.", bold_prefix="Asynchrony: ")

add_body("II. TypeScript/React Conventions")
add_bullet("Functional components exclusively with hooks (useState, useEffect, useContext, useReducer).", bold_prefix="Components: ")
add_bullet("TypeScript interfaces for all API types and component props. Strict mode enabled.", bold_prefix="Type Safety: ")
add_bullet("TailwindCSS utility classes with Radix UI accessible primitives.", bold_prefix="Styling: ")
add_bullet("Feature-domain folders (health/, latency/, tryit/) with shared UI in components/ui/.", bold_prefix="Organization: ")

add_body("III. Architecture Patterns")
add_bullet("AgentFactory.create() for backend-agnostic agent instantiation.", bold_prefix="Factory Pattern: ")
add_bullet("Pipecat processors for streaming voice processing.", bold_prefix="Pipeline Pattern: ")
add_bullet("DashboardContext with useReducer for centralized state.", bold_prefix="Context + Reducer: ")
add_bullet("All ML models loaded on first use via load_model() methods.", bold_prefix="Lazy Loading: ")

add_heading2("6.2 Project Coding")

add_body("1. src/api/main.py \u2014 FastAPI Application")

add_code_block(
    '"""\nFastAPI Application for Voice Assistant\n'
    'Main API endpoints + Dashboard for the hyper-localized multilingual voice assistant.\n"""\n\n'
    'import logging\nimport time\nfrom typing import Optional\nfrom pathlib import Path\nimport tempfile\n\n'
    'from dotenv import load_dotenv\nload_dotenv(Path(__file__).parent.parent.parent / ".env")\n\n'
    'from fastapi import FastAPI, UploadFile, File, Form, HTTPException, WebSocket\n'
    'from fastapi.middleware.cors import CORSMiddleware\n'
    'from fastapi.staticfiles import StaticFiles\nfrom fastapi.responses import FileResponse\n'
    'from pydantic import BaseModel\n\nlogger = logging.getLogger(__name__)\n\n'
    'STATIC_DIR = Path(__file__).parent / "static"\n\n'
    '# Initialize FastAPI app\napp = FastAPI(\n'
    '    title="Hyper-Localized Voice Assistant API",\n'
    '    description="Multilingual voice assistant with edge computing and agentic workflows",\n'
    '    version="0.2.0",\n)\n\n'
    '# CORS middleware\napp.add_middleware(\n    CORSMiddleware,\n'
    '    allow_origins=["*"],\n    allow_credentials=True,\n'
    '    allow_methods=["*"],\n    allow_headers=["*"],\n)\n\n'
    'class TranscriptionResponse(BaseModel):\n    text: str\n    language: str\n'
    '    language_name: str\n    confidence: Optional[float] = None\n\n'
    'class ProcessResponse(BaseModel):\n    transcription: str\n    intent: str\n'
    '    response_text: str\n    language: str\n\n'
    'class HealthResponse(BaseModel):\n    status: str\n    version: str\n'
    '    components: Optional[dict] = None',
    "The main entry point of the FastAPI backend. Initializes CORS middleware, defines Pydantic response models, "
    "and sets up REST endpoints and WebSocket connections."
)

add_body("2. src/agents/orchestrator.py \u2014 LangGraph Agent Orchestrator")

add_code_block(
    '"""\nLangGraph Agent Orchestrator\n'
    'Coordinates multiple agents using stateful graph-based workflows.\n"""\n\n'
    'import logging\nimport re\nimport time\n'
    'from typing import Dict, List, Any, Optional, TypedDict\nfrom enum import Enum\n\n'
    'logger = logging.getLogger(__name__)\n_agent_llm = None\n\n'
    'def _get_agent_llm():\n    """Get or create the shared LLM instance."""\n'
    '    global _agent_llm\n    if _agent_llm is None:\n        try:\n'
    '            from ..llm.qwen import get_llm\n'
    '            _agent_llm = get_llm(engine="qwen3", model_size="1.7b", device="cpu")\n'
    '        except Exception:\n            from ..llm.qwen import get_llm\n'
    '            _agent_llm = get_llm(engine="qwen", model_size="0.5b", device="cpu")\n'
    '    return _agent_llm\n\n'
    'class AgentState(TypedDict):\n    user_input: str\n    language: str\n'
    '    intent: Optional[str]\n    entities: List[Dict[str, Any]]\n'
    '    agent_outputs: Dict[str, Any]\n    final_response: Optional[str]\n    error: Optional[str]\n\n'
    'class IntentType(Enum):\n    INFORMATION_QUERY = "information_query"\n'
    '    TASK_MANAGEMENT = "task_management"\n    SMART_HOME = "smart_home"\n'
    '    GENERAL_CHAT = "general_chat"\n    UNKNOWN = "unknown"\n\n'
    'AGENT_PROMPTS = {\n    "info_agent": "You are an information assistant...",\n'
    '    "task_agent": "You are a task management assistant...",\n'
    '    "chat_agent": "You are a friendly conversational assistant...",\n}',
    "Core agent orchestrator using LangGraph stateful graph workflows. Defines agent state, "
    "intent classification types, and system prompts for each specialized agent."
)

add_body("3. src/db/models.py \u2014 Database Models")

add_code_block(
    '@dataclass\nclass Message:\n    """A single message in a conversation."""\n'
    '    id: str = ""\n    conversation_id: str = ""\n    role: str = ""\n'
    '    content: str = ""\n    language: str = "en"\n    timestamp: str = ""\n'
    '    metadata: Dict[str, Any] = field(default_factory=dict)\n\n'
    '@dataclass\nclass Conversation:\n    """A conversation session with multiple messages."""\n'
    '    id: str = ""\n    session_id: str = ""\n    language: str = "ml"\n'
    '    messages: List[Message] = field(default_factory=list)\n'
    '    asr_time_ms: float = 0\n    translation_time_ms: float = 0\n'
    '    llm_time_ms: float = 0\n    tts_time_ms: float = 0\n\n'
    '@dataclass\nclass Task:\n    """A user task created by the task agent."""\n'
    '    id: str = ""\n    title: str = ""\n    status: str = "pending"\n'
    '    created_at: str = ""\n    source_input: str = ""\n\n'
    '@dataclass\nclass SystemMetric:\n    """A system performance metric snapshot."""\n'
    '    id: str = ""\n    component: str = ""\n    metric_name: str = ""\n'
    '    value: float = 0.0',
    "Four core data models using Python dataclasses: Message, Conversation (with pipeline timing), "
    "Task (voice-created reminders), and SystemMetric (performance monitoring)."
)

add_body("4. frontend/src/App.tsx \u2014 React Dashboard")

add_code_block(
    'import { DashboardProvider } from "./context/DashboardContext";\n'
    'import { Header } from "./components/layout/Header";\n'
    'import { DashboardGrid, FullWidth } from "./components/layout/DashboardGrid";\n'
    'import { StatsRow } from "./components/stats/StatsRow";\n'
    'import { ComponentHealth } from "./components/health/ComponentHealth";\n'
    'import { LatencyBreakdown } from "./components/latency/LatencyBreakdown";\n'
    'import { TryItPanel } from "./components/tryit/TryItPanel";\n'
    'import { ConversationLog } from "./components/conversations/ConversationLog";\n'
    'import { TaskPanel } from "./components/tasks/TaskPanel";\n'
    'import { DashboardDataLoader } from "./DashboardDataLoader";\n\n'
    'export function App() {\n  return (\n    <DashboardProvider>\n'
    '      <DashboardDataLoader>\n        <Header />\n        <DashboardGrid>\n'
    '          <StatsRow />\n          <ComponentHealth />\n'
    '          <LatencyBreakdown />\n          <TryItPanel />\n'
    '          <TaskPanel />\n          <FullWidth>\n'
    '            <ConversationLog />\n          </FullWidth>\n'
    '        </DashboardGrid>\n      </DashboardDataLoader>\n'
    '    </DashboardProvider>\n  );\n}',
    "Main React component structuring the real-time monitoring dashboard with Provider pattern "
    "for state management and responsive grid layout."
)


# ══════════════════════════════════════════════════════════════
# CHAPTER 7 — TESTING
# ══════════════════════════════════════════════════════════════

add_heading1("7. TESTING")

add_heading2("7.1 Unit Testing \u2013 Test Cases")

add_table(
    ["Test Case ID", "Module", "Test Scenario", "Input", "Expected Output", "Actual Output", "Status"],
    [
        ["UTC-001", "whisper_asr.py", "Init Whisper ASR tiny", 'WhisperASR(model_size="tiny", device="cpu")', 'model_size="tiny", model=None', "As expected", "Pass"],
        ["UTC-002", "whisper_asr.py", "Supported Indian languages", "asr.SUPPORTED_INDIAN_LANGUAGES", 'Contains "hi", "ml"', "As expected", "Pass"],
        ["UTC-003", "pingala_asr.py", "Init Pingala ASR", 'PingalaASR(device="cpu")', "device='cpu', transcriber=None", "As expected", "Pass"],
        ["UTC-004", "pingala_asr.py", "Verify 11 languages", "len(SUPPORTED_INDIAN_LANGUAGES)", "11 languages", "As expected", "Pass"],
        ["UTC-005", "indic_tts.py", "Init Indic TTS", 'IndicTTS(default_language="hi")', 'default_language="hi"', "As expected", "Pass"],
        ["UTC-006", "indic_tts.py", "TTS model variants", "IndicTTS.MODEL_VARIANTS", '"mini" and "full"', "As expected", "Pass"],
        ["UTC-007", "tts_engine.py", "Supported emotions", "tts.get_supported_emotions()", '"neutral", "happy"', "As expected", "Pass"],
        ["UTC-008", "main.py (API)", "Health check", "GET /health", 'status="healthy", v0.2.0', "As expected", "Pass"],
        ["UTC-009", "main.py (API)", "Health components", "GET /health", "asr, tts, orchestrator, db", "As expected", "Pass"],
        ["UTC-010", "main.py (API)", "Dashboard stats", "GET /api/dashboard/stats", "conversations, latency, langs", "As expected", "Pass"],
        ["UTC-011", "main.py (API)", "Conversations list", "GET /api/dashboard/conversations", "conversations[], total", "As expected", "Pass"],
        ["UTC-012", "main.py (API)", "Not found conversation", "GET .../nonexistent", "HTTP 404 Not Found", "As expected", "Pass"],
    ],
)


# ══════════════════════════════════════════════════════════════
# CHAPTER 8 — IMPLEMENTATION
# ══════════════════════════════════════════════════════════════

add_heading1("8. IMPLEMENTATION")

add_body(
    "The implementation of the Hyper-Localized Multilingual Voice Assistant follows an agile, modular approach, "
    "focusing on parallel development of the AI pipeline, agent system, frontend dashboard, and edge deployment."
)

sections_8 = [
    ("I. Infrastructure Setup and Backend Development",
     "Set up Python 3.10+ virtual environment with PyTorch, Transformers, and Accelerate. Implemented SQLite "
     "database with WAL mode for concurrent access. Created FastAPI application with CORS middleware, static "
     "file serving, and Pydantic response models. Developed the modular AI pipeline connecting ASR \u2192 "
     "Translation \u2192 LLM \u2192 TTS with per-component timing metrics."),
    ("II. AI/ML Pipeline Development",
     "Integrated OpenAI Whisper for multilingual speech recognition (99+ languages). Implemented IndicTrans2 "
     "from AI4Bharat for bidirectional ML\u2194EN translation using the distributed variant (~211M params). "
     "Deployed Qwen 2.5 Instruct models via both Transformers and llama.cpp backends for GPU and edge inference. "
     "Added MMS-TTS for offline synthesis and an audio denoising pipeline using noisereduce."),
    ("III. Agent System Development",
     "Designed a LangGraph orchestrator with stateful graph-based workflows. Implemented four specialized "
     "agents: Info Agent, Task Agent, Chat Agent, and Smart Home Agent. Added LLM-based intent classification "
     "with keyword fallback. Created AgentFactory for backend-agnostic agent creation (LangGraph or ClawGo)."),
    ("IV. Frontend and Dashboard Development",
     "Built a React 18 + TypeScript SPA with Vite. Implemented TailwindCSS styling with Radix UI components. "
     "Added WebSocket streaming for real-time conversation updates, voice visualizer with waveform display, "
     "and an interactive \u201cTry It\u201d panel with microphone recording capability."),
    ("V. Testing and Quality Assurance",
     "Developed comprehensive test suite: unit tests for ASR, TTS, LLM, Translation initialization; API "
     "integration tests with FastAPI TestClient; database operation tests; cache functionality tests; and "
     "optimization tests. All tests run with pytest tests/ -v."),
    ("VI. Edge Deployment and Docker",
     "Created multi-stage Dockerfile with model caching. Docker Compose includes 2 GB memory reservation "
     "with 4 GB maximum, persistent volumes for models/data, and health checks. Raspberry Pi 5 client "
     "supports audio I/O, wake word detection, and model quantization (Q8_0) for edge inference."),
]

for title, content in sections_8:
    add_heading2(title)
    add_body(content)


# ══════════════════════════════════════════════════════════════
# CHAPTER 9 — BUSINESS PERSPECTIVE
# ══════════════════════════════════════════════════════════════

add_heading1("9. BUSINESS PERSPECTIVE")

add_body(
    "The Hyper-Localized Multilingual Voice Assistant is a strategic solution addressing a massive market "
    "gap in voice-first digital services for India\u2019s linguistically diverse population."
)

biz_sections = [
    ("I. Digital Inclusion and Accessibility", [
        ("Bridging the Digital Divide: ", "With only 10% of India\u2019s 1.4 billion population comfortable with English, this system opens digital services to the remaining 90% via native voice interaction."),
        ("Voice-First in Rural Areas: ", "In rural India, a voice interface eliminates the need for text-based interaction, making digital services accessible to a far wider population."),
        ("Reducing Literacy Barriers: ", "Voice commands and audio responses enable illiterate and semi-literate users to access digital services."),
    ]),
    ("II. Cost Efficiency", [
        ("Open-Source AI Stack: ", "Whisper, IndicTrans2, Qwen 2.5, MMS-TTS are all open-source, eliminating licensing fees."),
        ("Edge Deployment Economics: ", "A Raspberry Pi 5 (~$80) runs the entire system, eliminating recurring cloud costs."),
        ("Reduced Operational Costs: ", "Local processing eliminates bandwidth costs and cloud server provisioning."),
    ]),
    ("III. Privacy and Compliance", [
        ("Data Sovereignty: ", "All audio processing happens locally. No data leaves the local network."),
        ("Regulatory Compliance: ", "Inherently compliant with India\u2019s DPDPA 2023 as no personal data is transmitted externally."),
    ]),
    ("IV. Scalability", [
        ("Modular Architecture: ", "Adding a new language requires only a configuration change in settings.yaml."),
        ("Docker Deployment: ", "Single-command deployment via docker-compose up enables rapid scaling."),
        ("Market Size: ", "India\u2019s voice assistant market is projected at $1.6B by 2026; regional languages are the primary growth driver."),
    ]),
]

for section_title, items in biz_sections:
    add_heading2(section_title)
    for bold_t, rest_t in items:
        add_body_bold_start(bold_t, rest_t)


# ══════════════════════════════════════════════════════════════
# CHAPTER 10 — FUTURE SCOPE
# ══════════════════════════════════════════════════════════════

add_heading1("10. FUTURE SCOPE OR ENHANCEMENTS")

add_body(
    "The successful implementation lays the groundwork for strategic enhancements that expand functionality, "
    "improve user experience, and increase operational intelligence."
)

add_heading2("I. Advanced Features")
add_bullet("Implement Retrieval-Augmented Generation (RAG) for multi-turn conversation context.", bold_prefix="Conversation Memory with RAG: ")
add_bullet("Synthesize responses in a familiar voice using 5-second voice samples.", bold_prefix="Voice Cloning for Personalized TTS: ")
add_bullet("Custom deep learning model for natural \u201cHey Assistant\u201d activation.", bold_prefix="Deep Learning Wake Word: ")

add_heading2("II. Model Improvements")
add_bullet("Fine-tune IndicTrans2 on healthcare, agriculture, and education vocabulary.", bold_prefix="Domain-Specific Fine-Tuning: ")
add_bullet("Train custom ASR for specific Indian dialects (e.g., Malabar vs. Travancore Malayalam).", bold_prefix="Dialect-Specific ASR: ")
add_bullet("Real-time chunked audio processing for reduced perceived latency.", bold_prefix="Streaming ASR: ")

add_heading2("III. Platform Expansion")
add_bullet("React Native mobile app communicating with edge device over local Wi-Fi.", bold_prefix="Mobile Application: ")
add_bullet("Process voice messages through India\u2019s most popular messaging platform (500M+ users).", bold_prefix="WhatsApp Business API: ")
add_bullet("Synchronize across multiple edge nodes in a facility.", bold_prefix="Multi-Device Sync: ")


# ══════════════════════════════════════════════════════════════
# CHAPTER 11 — APPENDIX
# ══════════════════════════════════════════════════════════════

add_heading1("11. APPENDIX")

add_heading4("11.1 Git History [Sprint Wise]")

git_sprints = [
    ("Sprint 1 \u2013 Project Initialization & Core Pipeline", [
        "Set up project structure, CLAUDE.md development guidance, and Python environment.",
        "Implemented core ASR (Whisper), TTS (MMS), LLM (Qwen), and Translation (IndicTrans2) modules.",
        "Created initial FastAPI server with REST endpoints.",
        "Established the end-to-end pipeline: Audio \u2192 ASR \u2192 Translation \u2192 LLM \u2192 TTS \u2192 Audio.",
    ]),
    ("Sprint 2 \u2013 Frontend Dashboard & UI", [
        "Created React frontend dashboard with Vite + TypeScript scaffolding.",
        "Removed node_modules from tracking and added root README.",
        "Migrated to TailwindCSS, added voice visualizer, updated backend pipeline.",
        "Implemented StatsRow, ComponentHealth, ConversationLog components.",
    ]),
    ("Sprint 3 \u2013 Docker, Optimization & Documentation", [
        "Added multi-stage Docker setup and docker-compose.yml with memory limits.",
        "Implemented LRU cache with TTL for response caching.",
        "Added model optimization utilities (quantization, memory profiling).",
        "Created benchmarking script and sprint documentation.",
    ]),
    ("Sprint 4 \u2013 Testing, EdgeX Bridge & Agent Enhancement", [
        "Added comprehensive test suite: test_basic, test_api, test_database, test_cache.",
        "Implemented XiaoZhi/EdgeX protocol bridge for ESP32 communication.",
        "Added LLM-based intent classification in the agent orchestrator.",
    ]),
    ("Sprint 5 \u2013 Agent Routing, Seed Data & Demo", [
        "Added agent routing display in the dashboard UI.",
        "Implemented seed_demo_data.py for sample conversations.",
        "Switched to offline-only models for reliable demonstration.",
    ]),
    ("Sprint 6 \u2013 Task Management, Web Search & Audio", [
        "Added task management with Task model, API endpoints, TaskPanel UI.",
        "Implemented web search integration for real-time information.",
        "Added audio denoising pipeline using noisereduce.",
        "Final UI enhancements and polish across all dashboard components.",
    ]),
]

for sprint_title, items in git_sprints:
    add_body_bold_start(sprint_title, "")
    for item in items:
        add_bullet(item)

add_heading3("11.2 Screenshots of Forms, Reports, Diagrams")

add_body("Dashboard Overview — Stats, Component Health, and Latency Breakdown")
add_figure("screenshot_dashboard_full.png", "Fig 11.1 \u2014 Dashboard: Statistics Row, Component Health, and Average Latency Breakdown", Inches(5.8))

add_body("Try It Panel and Conversation Log")
add_figure("screenshot_dashboard_mid.png", "Fig 11.2 \u2014 Try It Panel (Voice Input) and Recent Conversations with Multilingual Queries", Inches(5.8))

add_body("Recent Conversations — Multilingual Query History")
add_figure("screenshot_dashboard_bottom.png", "Fig 11.3 \u2014 Conversation Log with Language Tags, Intent Classification, and Timestamps", Inches(5.8))

add_body("Complete Dashboard View")
add_figure("screenshot_dashboard_fullpage.png", "Fig 11.4 \u2014 Full-Page Dashboard: End-to-End View of All Monitoring Panels", Inches(5.5))

add_body("Responsive Layout (Narrow Viewport)")
add_figure("screenshot_dashboard_narrow.png", "Fig 11.5 \u2014 Dashboard Responsive Layout at 800px Width", Inches(4))

add_heading3("11.3 Burndown Chart")

add_body(
    "The burndown chart tracks remaining effort across six sprints. The actual line closely follows the "
    "ideal, with minor deviations in Sprint 5 (reduced scope as system approached demo-readiness)."
)

add_figure("burndown.png", "Fig 11.6 \u2014 Burndown Chart (repeated for reference)", Inches(5))


# ══════════════════════════════════════════════════════════════
# CHAPTER 12 — REFERENCES
# ══════════════════════════════════════════════════════════════

add_heading1("12. References / Bibliography")

add_body("Books and Articles")
refs_books = [
    "Russell, S. & Norvig, P. (2021). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.",
    "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press.",
    "Jurafsky, D. & Martin, J.H. (2023). Speech and Language Processing (3rd ed. draft). Stanford University.",
]
for r in refs_books:
    p = add_body(r)

add_body("Research Papers")
refs_papers = [
    "Radford, A., et al. (2023). Robust Speech Recognition via Large-Scale Weak Supervision. ICML 2023. OpenAI.",
    "Costa-juss\u00e0, M.R., et al. (2022). No Language Left Behind: Scaling Human-Centered Machine Translation. Meta AI.",
    "Pratap, V., et al. (2023). Scaling Speech Technology to 1,000+ Languages. Meta AI.",
    "AI4Bharat, Gala, J., et al. (2023). IndicTrans2: High-Quality MT Models for 22 Indian Languages. TACL.",
]
for r in refs_papers:
    add_body(r)

add_body("Websites")
refs_web = [
    "https://huggingface.co \u2014 Model hosting and inference APIs.",
    "https://pytorch.org \u2014 PyTorch deep learning framework.",
    "https://fastapi.tiangolo.com \u2014 FastAPI web framework.",
    "https://langchain-ai.github.io/langgraph/ \u2014 LangGraph agent framework.",
    "https://ai4bharat.org \u2014 Indian language AI research.",
]
for r in refs_web:
    add_body(r)

add_body("Technical Documentation")
refs_tech = [
    "OpenAI Whisper \u2014 https://github.com/openai/whisper",
    "Meta MMS \u2014 https://github.com/facebookresearch/fairseq/tree/main/examples/mms",
    "Qwen 2.5 \u2014 https://github.com/QwenLM/Qwen2.5",
    "IndicTrans2 \u2014 https://github.com/AI4Bharat/IndicTrans2",
    "React Official Docs \u2014 https://react.dev",
]
for r in refs_tech:
    add_body(r)

add_body("Standards and Guidelines")
add_body("ISO/IEC 25010:2011 \u2014 Systems and software quality models.")
add_body("W3C WCAG 2.1 \u2014 Web Content Accessibility Guidelines.")
add_body("India\u2019s Digital Personal Data Protection Act (DPDPA) 2023.")


# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════

doc.save(str(OUT))
print(f"\nReport saved to: {OUT}")
print(f"Total paragraphs: ~{len(doc.paragraphs)}")
